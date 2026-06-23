# =============================================================
# DOCX TEXT EXTRACTION LIBRARIES - COMPARISON & EXAMPLES
# =============================================================


# -------------------------------------------------------------
# 1. python-docx
#    Official, most popular library
#    Extracts paragraphs, tables, headers, footers
#    Can read styles, fonts, formatting metadata
#    Cannot handle password protected files
#    Best for: clean structured docx files
# -------------------------------------------------------------
# pip install python-docx

from docx import Document
def extract_python_docx(path):
    doc  = Document(path)
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text.strip() + " "
            text += "\n"
    return text.strip()
text = extract_python_docx("file.docx")
print(text)


# -------------------------------------------------------------
# 2. docx2txt
#    Extremely simple, one line extraction
#    Extracts text and images
#    No control over what gets extracted
#    Best for: quick dirty extraction when you just need raw text
# -------------------------------------------------------------
# pip install docx2txt

import docx2txt
def extract_docx2txt(path):
    text = docx2txt.process(path)
    return text.strip()
text = extract_docx2txt("file.docx")
print(text)
# extract text and save images to a folder
text = docx2txt.process("file.docx", "SYSTEM/images/")
print(text)


# -------------------------------------------------------------
# 3. mammoth
#    Converts docx to clean HTML or markdown
#    Best at preserving document structure
#    Great for documents with heavy formatting
#    Best for: when you need headings, bold, lists preserved
# -------------------------------------------------------------
# pip install mammoth

import mammoth
def extract_mammoth_html(path):
    with open(path, "rb") as f:
        result = mammoth.convert_to_html(f)
    return result.value
def extract_mammoth_markdown(path):
    with open(path, "rb") as f:
        result = mammoth.convert_to_markdown(f)
    return result.value
def extract_mammoth_text(path):
    with open(path, "rb") as f:
        result = mammoth.extract_raw_text(f)
    return result.value.strip()
html     = extract_mammoth_html("file.docx")
markdown = extract_mammoth_markdown("file.docx")
text     = extract_mammoth_text("file.docx")
print(html)
print(markdown)
print(text)


# -------------------------------------------------------------
# 4. textract
#    Extracts text from many formats (docx, pdf, xlsx, pptx, csv)
#    Good for pipelines handling multiple file types
#    Heavy dependency, harder to install on Windows
#    Best for: multi-format pipelines
# -------------------------------------------------------------
# pip install textract

import textract
def extract_textract(path):
    text = textract.process(path)
    return text.decode("utf-8").strip()
# works the same way for any file type
docx_text = extract_textract("file.docx")
pdf_text  = extract_textract("file.pdf")
xlsx_text = extract_textract("file.xlsx")
pptx_text = extract_textract("file.pptx")
print(docx_text)


# -------------------------------------------------------------
# 5. Apache Tika (via tika-python)
#    Runs a Java server in background
#    Handles 1000+ file formats
#    Very powerful but requires Java installed
#    Best for: enterprise pipelines with many file types
# -------------------------------------------------------------
# pip install tika
# requires Java installed: https://www.java.com/en/download/

from tika import parser
def extract_tika(path):
    parsed = parser.from_file(path)
    text   = parsed.get("content", "") or ""
    return text.strip()
# works the same way for any file type
docx_text = extract_tika("file.docx")
pdf_text  = extract_tika("file.pdf")
xlsx_text = extract_tika("file.xlsx")
print(docx_text)


# =============================================================
# QUICK COMPARISON
# =============================================================
#
# Simple docx files          -> python-docx
# Heavy formatted documents  -> mammoth
# Multiple file types        -> textract
# Enterprise / many formats  -> tika
# Just need raw text fast    -> docx2txt
#
# Install all at once:
# pip install python-docx docx2txt mammoth textract tika
# =============================================================



# Library       | Tables | Images | Formatting | Multi-format | Windows Easy | Speed
# --------------|--------|--------|------------|--------------|--------------|--------
# python-docx   | Yes    | No     | Yes        | No           | Yes          | Fast
# docx2txt      | Partial| Yes    | No         | No           | Yes          | Very Fast
# mammoth       | Yes    | No     | HTML/MD    | No           | Yes          | Fast
# textract      | Yes    | No     | No         | Yes          | Hard         | Medium
# tika          | Yes    | Yes    | Yes        | Yes          | Needs Java   | Slow